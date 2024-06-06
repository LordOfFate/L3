from docx import Document
varname = '2.docx'
doc = Document(varname)
out = []
for i in doc.paragraphs:
    huh = i.text
    lol = 0
    for j in huh:
        if j == " ":
            lol += 1
    if lol%2 == 0:
        out.append(huh + " YESYESYES" + "\n")
    else: 
        out.append(huh + " NONONO" + "\n")
    lol = 0
for i in out:
    print(i)
doc = Document()
for i in out:
    doc.add_paragraph(i)
doc.save('d.docx')